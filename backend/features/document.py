"""
문서 관련 실제 기능 코드.
기본 저장 형식: .docx (Word)
"""
from pathlib import Path
from storage.database import get_connection

# workspace 경로 (main.py와 동일 기준)
WORKSPACE_DIR = Path(__file__).parent.parent.parent / "workspace"


def _doc_path(title: str) -> str:
    """제목을 파일명으로 변환해 workspace 경로 반환"""
    safe_title = "".join(c for c in title if c not in r'\/:*?"<>|').strip()
    if not safe_title:
        safe_title = "제목없음"
    return str(WORKSPACE_DIR / f"{safe_title}.docx")


def create_document() -> dict:
    """새 문서 생성 — DB에 등록 후 빈 .docx 파일 생성"""
    conn = get_connection()
    cur = conn.execute(
        "INSERT INTO documents (title, content) VALUES (?, ?)",
        ("제목 없음", "")
    )
    conn.commit()
    doc_id = cur.lastrowid
    conn.close()

    # 빈 .docx 파일 생성
    _create_docx_file("제목 없음", "")

    return get_document(doc_id)


def get_document(doc_id: int) -> dict | None:
    """문서 단건 조회"""
    conn = get_connection()
    row = conn.execute(
        "SELECT id, title, content, mode, created_at, updated_at "
        "FROM documents WHERE id = ?", (doc_id,)
    ).fetchone()
    conn.close()
    if not row:
        return None
    return dict(row)


def save_document(doc_id: int, title: str, content: str, mode: str) -> dict:
    """문서 저장 — DB 업데이트 + .docx 파일 저장"""
    conn = get_connection()

    # 이전 제목 조회 (파일명 변경 감지)
    old = conn.execute("SELECT title FROM documents WHERE id=?", (doc_id,)).fetchone()
    old_title = old["title"] if old else title

    conn.execute(
        """UPDATE documents
           SET title=?, content=?, mode=?,
               updated_at=datetime('now','localtime')
           WHERE id=?""",
        (title, content, mode, doc_id)
    )
    conn.commit()
    conn.close()

    # 제목이 바뀌면 이전 파일 삭제
    if old_title != title:
        old_path = Path(_doc_path(old_title))
        if old_path.exists():
            old_path.unlink()

    # .docx 저장
    _create_docx_file(title, content)

    return get_document(doc_id)


def autosave_document(doc_id: int, content: str) -> bool:
    """자동 저장 — DB + .docx 파일 동시 저장"""
    conn = get_connection()
    row = conn.execute("SELECT title FROM documents WHERE id=?", (doc_id,)).fetchone()
    title = row["title"] if row else "제목없음"

    conn.execute(
        """UPDATE documents
           SET content=?, updated_at=datetime('now','localtime')
           WHERE id=?""",
        (content, doc_id)
    )
    conn.commit()
    conn.close()

    # .docx 파일도 갱신
    _create_docx_file(title, content)
    return True


def list_documents() -> list[dict]:
    """문서 목록 (최신순)"""
    conn = get_connection()
    rows = conn.execute(
        "SELECT id, title, mode, updated_at FROM documents ORDER BY updated_at DESC"
    ).fetchall()
    conn.close()
    return [dict(r) for r in rows]


def delete_document(doc_id: int) -> bool:
    """문서 삭제 — DB + .docx 파일 삭제"""
    conn = get_connection()
    row = conn.execute("SELECT title FROM documents WHERE id=?", (doc_id,)).fetchone()
    if row:
        path = Path(_doc_path(row["title"]))
        if path.exists():
            path.unlink()
    conn.execute("DELETE FROM documents WHERE id=?", (doc_id,))
    conn.commit()
    conn.close()
    return True


def count_text(content: str) -> dict:
    """글자 수 / 단어 수 / 줄 수 계산"""
    no_space = content.replace(" ", "").replace("\n", "").replace("\t", "")
    words = [w for w in content.split() if w]
    lines = content.split("\n")
    return {
        "chars":          len(content),
        "chars_no_space": len(no_space),
        "words":          len(words),
        "lines":          len(lines),
    }


def _create_docx_file(title: str, content: str):
    """텍스트를 .docx 파일로 workspace에 저장"""
    try:
        from docx import Document
        from docx.shared import Pt, Cm

        doc = Document()
        section = doc.sections[0]
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(3.0)

        for line in content.split("\n"):
            stripped = line.strip()
            if stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            else:
                p = doc.add_paragraph(line)
                p.style.font.size = Pt(11)

        WORKSPACE_DIR.mkdir(exist_ok=True)
        doc.save(_doc_path(title))

    except ImportError:
        # python-docx 미설치 시 .txt로 대체 저장
        path = WORKSPACE_DIR / f"{title}.txt"
        path.write_text(content, encoding="utf-8")