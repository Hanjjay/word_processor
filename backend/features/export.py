"""
파일 내보내기 기능 코드.
docx / pdf 지원.
"""


def to_docx(content: str, output_path: str) -> str:
    """
    텍스트를 .docx 파일로 저장.
    pip install python-docx 필요.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm
    except ImportError:
        raise ImportError("python-docx 미설치. 터미널에서: pip install python-docx")

    doc = Document()

    # A4 여백 설정
    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(3.0)

    for line in content.split("\n"):
        stripped = line.strip()
        # 마크다운 헤딩 자동 변환
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        else:
            p = doc.add_paragraph(line)
            p.style.font.size = Pt(11)

    doc.save(output_path)
    return output_path


def to_pdf(html_content: str, output_path: str) -> str:
    """
    HTML을 .pdf 파일로 저장.
    pip install weasyprint 필요.
    """
    try:
        from weasyprint import HTML, CSS
    except ImportError:
        raise ImportError("weasyprint 미설치. 터미널에서: pip install weasyprint")

    css = CSS(string="""
        @page { size: A4; margin: 2.5cm 3cm; }
        body {
            font-family: 'Malgun Gothic', 'NanumGothic', sans-serif;
            font-size: 11pt;
            line-height: 1.8;
            color: #1a1a1a;
        }
    """)
    HTML(string=html_content).write_pdf(output_path, stylesheets=[css])
    return output_path
