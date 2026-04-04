def export_to_docx(plain_text: str, output_path: str) -> None:
    """
    에디터의 텍스트를 .docx 파일로 저장합니다.
    pip install python-docx 필요.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx가 설치되지 않았습니다.\n터미널에서: pip install python-docx")

    doc = Document()

    # 페이지 여백 설정
    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(3)

    # 줄 단위로 삽입 (빈 줄은 단락 구분)
    lines = plain_text.split("\n")
    for line in lines:
        stripped = line.strip()

        # 마크다운 헤딩 변환
        if stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            p = doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            p = doc.add_heading(stripped[4:], level=3)
        else:
            p = doc.add_paragraph(line)
            p.style.font.size = Pt(11)

    doc.save(output_path)
