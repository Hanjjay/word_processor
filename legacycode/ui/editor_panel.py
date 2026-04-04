import os
from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QTextEdit, QScrollArea,
    QFileDialog, QMessageBox, QToolBar, QComboBox, QFrame
)
from PyQt6.QtGui import (
    QFont, QTextCharFormat, QTextCursor,
    QAction, QKeySequence, QPageSize, QPageLayout
)
from PyQt6.QtCore import Qt, pyqtSignal, QMarginsF
from PyQt6.QtPrintSupport import QPrinter

from core.document import Document
from core.snapshot import SnapshotManager
from export.to_docx import export_to_docx
from export.to_pdf import export_to_pdf


# A4 픽셀 크기 (96dpi 기준)
A4_W_PX = 794
A4_H_PX = 1123
MARGIN_PX = 80   # 상하좌우 여백


class A4Editor(QTextEdit):
    """
    A4 용지처럼 보이는 QTextEdit.
    흰 배경 + 고정 너비로 실제 문서 느낌을 줍니다.
    """
    def __init__(self):
        super().__init__()
        self.setFont(QFont("맑은 고딕", 12))
        self.setPlaceholderText("글을 시작하세요...")

        # A4 크기 고정
        self.setFixedWidth(A4_W_PX)
        self.setMinimumHeight(A4_H_PX)

        # 내부 여백 (마진)
        self.setViewportMargins(MARGIN_PX, MARGIN_PX, MARGIN_PX, MARGIN_PX)

        self.setStyleSheet("""
            QTextEdit {
                background-color: #ffffff;
                color: #1a1a1a;
                border: none;
                selection-background-color: #cce5ff;
                font-size: 12pt;
                line-height: 1.8;
            }
        """)

        # 자동 줄바꿈
        self.setLineWrapMode(QTextEdit.LineWrapMode.WidgetWidth)


class EditorPanel(QWidget):
    mode_changed = pyqtSignal(str)   # 상태 표시줄 모드 업데이트용

    def __init__(self):
        super().__init__()
        self.current_document = Document()
        self.snapshot_manager = SnapshotManager()
        self._build_ui()

    # ── UI 구성 ──────────────────────────────────────────

    def _build_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        # 에디터 툴바
        layout.addWidget(self._build_toolbar())

        # 스크롤 영역 (회색 배경 = 책상)
        scroll = QScrollArea()
        scroll.setWidgetResizable(False)
        scroll.setAlignment(Qt.AlignmentFlag.AlignHCenter)  # A4 용지 가운데 정렬
        scroll.setStyleSheet("""
            QScrollArea {
                background-color: #2d2d2d;
                border: none;
            }
            QScrollBar:vertical {
                background: #2d2d2d;
                width: 10px;
            }
            QScrollBar::handle:vertical {
                background: #555;
                border-radius: 5px;
                min-height: 20px;
            }
        """)

        # A4 용지 컨테이너 (그림자 효과용 래퍼)
        paper_wrap = QWidget()
        paper_wrap.setStyleSheet("background-color: #2d2d2d; padding: 32px 0;")
        wrap_layout = QVBoxLayout(paper_wrap)
        wrap_layout.setAlignment(Qt.AlignmentFlag.AlignHCenter)
        wrap_layout.setContentsMargins(0, 32, 0, 32)

        # A4 용지 (흰 박스 + 그림자)
        paper_frame = QFrame()
        paper_frame.setStyleSheet("""
            QFrame {
                background-color: #ffffff;
                border: 1px solid #cccccc;
                border-radius: 2px;
            }
        """)
        # box-shadow 효과는 QGraphicsDropShadowEffect로 구현
        from PyQt6.QtWidgets import QGraphicsDropShadowEffect
        from PyQt6.QtGui import QColor
        shadow = QGraphicsDropShadowEffect()
        shadow.setBlurRadius(20)
        shadow.setOffset(0, 4)
        shadow.setColor(QColor(0, 0, 0, 80))
        paper_frame.setGraphicsEffect(shadow)

        paper_layout = QVBoxLayout(paper_frame)
        paper_layout.setContentsMargins(0, 0, 0, 0)

        # 실제 에디터
        self.editor = A4Editor()
        paper_layout.addWidget(self.editor)

        paper_frame.setFixedWidth(A4_W_PX)
        wrap_layout.addWidget(paper_frame)

        scroll.setWidget(paper_wrap)
        layout.addWidget(scroll)

    def _build_toolbar(self):
        toolbar = QToolBar()
        toolbar.setMovable(False)
        toolbar.setStyleSheet("""
            QToolBar {
                background-color: #252526;
                border-bottom: 1px solid #3c3c3c;
                padding: 3px 6px;
                spacing: 2px;
            }
            QToolButton {
                color: #cccccc;
                background: transparent;
                border: none;
                border-radius: 3px;
                padding: 3px 8px;
                font-size: 13px;
            }
            QToolButton:hover { background-color: #3c3c3c; }
            QComboBox {
                background-color: #3c3c3c;
                color: #cccccc;
                border: 1px solid #555;
                border-radius: 3px;
                padding: 2px 6px;
                min-width: 90px;
            }
            QComboBox::drop-down { border: none; }
            QComboBox QAbstractItemView {
                background-color: #252526;
                color: #cccccc;
                selection-background-color: #094771;
            }
        """)

        # 모드 선택
        self.mode_selector = QComboBox()
        self.mode_selector.addItems(["일반", "마크다운", "대본", "뮤지컬 가사"])
        self.mode_selector.currentTextChanged.connect(self._on_mode_changed)
        toolbar.addWidget(self.mode_selector)

        toolbar.addSeparator()

        # 서식 버튼
        for label, shortcut, slot in [
            ("B",  "Ctrl+B", self._bold),
            ("I",  "Ctrl+I", self._italic),
            ("U",  "Ctrl+U", self._underline),
        ]:
            a = QAction(label, self)
            a.setShortcut(QKeySequence(shortcut))
            a.triggered.connect(slot)
            toolbar.addAction(a)

        toolbar.addSeparator()

        # 글꼴 크기
        self.font_size = QComboBox()
        self.font_size.addItems(["10","11","12","13","14","16","18","20","24","28","32"])
        self.font_size.setCurrentText("12")
        self.font_size.setFixedWidth(60)
        self.font_size.currentTextChanged.connect(
            lambda s: self._apply_fmt(lambda f: f.setFontPointSize(float(s)))
        )
        toolbar.addWidget(self.font_size)

        return toolbar

    # ── 서식 ─────────────────────────────────────────────

    def _apply_fmt(self, fn):
        fmt = QTextCharFormat()
        fn(fmt)
        self.editor.textCursor().mergeCharFormat(fmt)

    def _bold(self):
        cur = self.editor.textCursor()
        w = QFont.Weight.Normal if cur.charFormat().fontWeight() == QFont.Weight.Bold else QFont.Weight.Bold
        self._apply_fmt(lambda f: f.setFontWeight(w))

    def _italic(self):
        cur = self.editor.textCursor()
        self._apply_fmt(lambda f: f.setFontItalic(not cur.charFormat().fontItalic()))

    def _underline(self):
        cur = self.editor.textCursor()
        self._apply_fmt(lambda f: f.setFontUnderline(not cur.charFormat().fontUnderline()))

    # ── 모드 ─────────────────────────────────────────────

    def _on_mode_changed(self, mode: str):
        self.mode_changed.emit(mode)
        fonts = {
            "일반":       ("맑은 고딕", 12, "#ffffff"),
            "마크다운":   ("Consolas",  12, "#1e1e2e"),
            "대본":       ("Courier New", 12, "#ffffff"),
            "뮤지컬 가사":("맑은 고딕", 13, "#fffdf0"),
        }
        family, size, bg = fonts.get(mode, ("맑은 고딕", 12, "#ffffff"))
        self.editor.setFont(QFont(family, size))
        self.editor.setStyleSheet(f"""
            QTextEdit {{
                background-color: {bg};
                color: #1a1a1a;
                border: none;
                font-size: {size}pt;
                line-height: 1.8;
            }}
        """)

    # ── 파일 작업 ─────────────────────────────────────────

    def new_document(self):
        if self.editor.document().isModified():
            r = QMessageBox.question(self, "새 문서", "저장하지 않은 내용이 있습니다. 계속할까요?")
            if r == QMessageBox.StandardButton.No:
                return
        self.editor.clear()
        self.current_document = Document()

    def open_document(self):
        path, _ = QFileDialog.getOpenFileName(
            self, "파일 열기", "",
            "텍스트/마크다운 (*.txt *.md);;Word (*.docx);;모든 파일 (*)"
        )
        if path:
            self.open_file(path)

    def open_file(self, path: str):
        """파일 트리에서 클릭했을 때도 사용"""
        ext = os.path.splitext(path)[1].lower()
        try:
            if ext in (".txt", ".md"):
                with open(path, "r", encoding="utf-8") as f:
                    self.editor.setPlainText(f.read())
                if ext == ".md":
                    self.mode_selector.setCurrentText("마크다운")
            elif ext == ".docx":
                from docx import Document as DX
                doc = DX(path)
                self.editor.setPlainText("\n".join(p.text for p in doc.paragraphs))
            else:
                with open(path, "r", encoding="utf-8", errors="replace") as f:
                    self.editor.setPlainText(f.read())
        except Exception as e:
            QMessageBox.warning(self, "오류", str(e))
        self.current_document.path = path

    def save_document(self):
        if not self.current_document.path:
            path, _ = QFileDialog.getSaveFileName(self, "저장", "", "텍스트 (*.txt);;마크다운 (*.md)")
            if not path:
                return
            self.current_document.path = path
        with open(self.current_document.path, "w", encoding="utf-8") as f:
            f.write(self.editor.toPlainText())
        self.editor.document().setModified(False)

    def export_docx(self):
        path, _ = QFileDialog.getSaveFileName(self, "Word 내보내기", "", "Word (*.docx)")
        if path:
            try:
                export_to_docx(self.editor.toPlainText(), path)
                QMessageBox.information(self, "완료", f"저장: {path}")
            except ImportError as e:
                QMessageBox.warning(self, "오류", str(e))

    def export_pdf(self):
        path, _ = QFileDialog.getSaveFileName(self, "PDF 내보내기", "", "PDF (*.pdf)")
        if path:
            try:
                export_to_pdf(self.editor.toHtml(), path)
                QMessageBox.information(self, "완료", f"저장: {path}")
            except ImportError as e:
                QMessageBox.warning(self, "오류", str(e))

    # ── 스냅샷 ───────────────────────────────────────────

    def take_snapshot(self):
        sid = self.snapshot_manager.take(self.editor.toPlainText())
        QMessageBox.information(self, "스냅샷", f"스냅샷 #{sid} 저장됨")

    def show_snapshots(self):
        snaps = self.snapshot_manager.list_all()
        if not snaps:
            QMessageBox.information(self, "스냅샷", "저장된 스냅샷이 없습니다.")
            return
        QMessageBox.information(self, "스냅샷 목록",
            "\n".join(f"#{s['id']}  {s['timestamp']}" for s in snaps))

    # ── 자동 저장용 ───────────────────────────────────────

    def get_content(self) -> str:
        return self.editor.toPlainText()

    def get_document_id(self) -> int:
        return self.current_document.doc_id
