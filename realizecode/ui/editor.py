import os
from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QScrollArea,
    QTextEdit, QToolBar, QComboBox,
    QFileDialog, QMessageBox, QGraphicsDropShadowEffect
)
from PyQt6.QtGui import (
    QFont, QTextCharFormat, QColor,
    QAction, QKeySequence
)
from PyQt6.QtCore import Qt, pyqtSignal


# ── A4 크기 (96 dpi 기준) ────────────────────────────────
A4_W  = 794    # px
A4_H  = 1123   # px
PAD   = 80     # 용지 내부 상하좌우 여백


class A4Paper(QTextEdit):
    """
    흰 A4 용지처럼 보이는 텍스트 에디터.
    고정 너비 794px, 내부 여백 80px.
    """
    def __init__(self):
        super().__init__()
        self.setFont(QFont("맑은 고딕", 11))
        self.setPlaceholderText("여기에 글을 쓰세요...")

        # 폭 고정 (A4)
        self.setFixedWidth(A4_W)
        self.setMinimumHeight(A4_H)

        # 용지 안쪽 여백
        self.setViewportMargins(PAD, PAD, PAD, PAD)

        # 줄바꿈: 위젯 너비 기준
        self.setLineWrapMode(QTextEdit.LineWrapMode.WidgetWidth)

        self.setStyleSheet("""
            QTextEdit {
                background-color: #ffffff;
                color: #1a1a1a;
                border: none;
                selection-background-color: #b3d4ff;
                font-size: 11pt;
                line-height: 1.8;
            }
        """)


class EditorPanel(QWidget):
    """
    툴바 + 스크롤 영역 + A4 용지로 구성된 에디터 패널.
    배경은 어두운 회색 (Word 다크 모드와 유사).
    """
    mode_changed = pyqtSignal(str)    # 상태 표시줄용
    saved        = pyqtSignal()       # 저장 완료 알림

    def __init__(self):
        super().__init__()
        self._current_path = ""
        self._build()

    # ── UI 구성 ──────────────────────────────────────────

    def _build(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        # 1. 툴바
        layout.addWidget(self._build_toolbar())

        # 2. 스크롤 영역 (다크 회색 = '책상')
        scroll = QScrollArea()
        scroll.setWidgetResizable(False)
        scroll.setAlignment(Qt.AlignmentFlag.AlignHCenter)
        scroll.setStyleSheet("""
            QScrollArea {
                background-color: #3c3c3c;
                border: none;
            }
            QScrollBar:vertical {
                background: #3c3c3c;
                width: 10px;
                margin: 0;
            }
            QScrollBar::handle:vertical {
                background: #666666;
                border-radius: 5px;
                min-height: 20px;
            }
            QScrollBar::add-line:vertical,
            QScrollBar::sub-line:vertical {
                height: 0;
            }
            QScrollBar:horizontal {
                background: #3c3c3c;
                height: 10px;
            }
            QScrollBar::handle:horizontal {
                background: #666666;
                border-radius: 5px;
            }
        """)

        # 3. 용지를 감싸는 컨테이너 (상하 여백용)
        container = QWidget()
        container.setStyleSheet("background-color: #3c3c3c;")
        c_layout = QVBoxLayout(container)
        c_layout.setAlignment(Qt.AlignmentFlag.AlignHCenter)
        c_layout.setContentsMargins(40, 40, 40, 40)
        c_layout.setSpacing(0)

        # 4. A4 용지 + 그림자
        self.text_edit = A4Paper()

        shadow = QGraphicsDropShadowEffect()
        shadow.setBlurRadius(24)
        shadow.setOffset(0, 3)
        shadow.setColor(QColor(0, 0, 0, 120))
        self.text_edit.setGraphicsEffect(shadow)

        c_layout.addWidget(self.text_edit)

        scroll.setWidget(container)
        layout.addWidget(scroll)

    # ── 툴바 ─────────────────────────────────────────────

    def _build_toolbar(self):
        tb = QToolBar()
        tb.setMovable(False)
        tb.setStyleSheet("""
            QToolBar {
                background-color: #2d2d2d;
                border-bottom: 1px solid #3c3c3c;
                padding: 2px 6px;
                spacing: 2px;
            }
            QToolButton {
                color: #cccccc;
                background: transparent;
                border: none;
                border-radius: 3px;
                padding: 3px 8px;
                font-size: 13px;
                min-width: 24px;
            }
            QToolButton:hover {
                background-color: #3c3c3c;
            }
            QComboBox {
                background-color: #3c3c3c;
                color: #cccccc;
                border: 1px solid #555555;
                border-radius: 3px;
                padding: 2px 6px;
                min-width: 80px;
                font-size: 12px;
            }
            QComboBox::drop-down { border: none; }
            QComboBox QAbstractItemView {
                background-color: #252526;
                color: #cccccc;
                selection-background-color: #094771;
                border: 1px solid #555;
            }
        """)

        # 글쓰기 모드
        self.mode_box = QComboBox()
        self.mode_box.addItems(["일반", "마크다운", "대본", "뮤지컬 가사"])
        self.mode_box.currentTextChanged.connect(self._change_mode)
        tb.addWidget(self.mode_box)

        tb.addSeparator()

        # 서식 버튼
        for text, shortcut, fn in [
            ("B",   "Ctrl+B", self._bold),
            ("I",   "Ctrl+I", self._italic),
            ("U",   "Ctrl+U", self._underline),
        ]:
            a = QAction(text, self)
            a.setShortcut(QKeySequence(shortcut))
            a.triggered.connect(fn)
            tb.addAction(a)

        tb.addSeparator()

        # 글꼴 크기
        self.size_box = QComboBox()
        self.size_box.addItems(["9","10","11","12","13","14","16","18","20","24","28"])
        self.size_box.setCurrentText("11")
        self.size_box.setFixedWidth(55)
        self.size_box.currentTextChanged.connect(self._change_size)
        tb.addWidget(self.size_box)

        return tb

    # ── 서식 ─────────────────────────────────────────────

    def _fmt(self, fn):
        """현재 선택 영역에 QTextCharFormat 적용"""
        fmt = QTextCharFormat()
        fn(fmt)
        self.text_edit.textCursor().mergeCharFormat(fmt)

    def _bold(self):
        cur = self.text_edit.textCursor()
        is_bold = cur.charFormat().fontWeight() == QFont.Weight.Bold
        self._fmt(lambda f: f.setFontWeight(
            QFont.Weight.Normal if is_bold else QFont.Weight.Bold
        ))

    def _italic(self):
        cur = self.text_edit.textCursor()
        self._fmt(lambda f: f.setFontItalic(not cur.charFormat().fontItalic()))

    def _underline(self):
        cur = self.text_edit.textCursor()
        self._fmt(lambda f: f.setFontUnderline(not cur.charFormat().fontUnderline()))

    def _change_size(self, size_str: str):
        try:
            self._fmt(lambda f: f.setFontPointSize(float(size_str)))
        except ValueError:
            pass

    # ── 모드 전환 ─────────────────────────────────────────

    def _change_mode(self, mode: str):
        self.mode_changed.emit(f"  {mode}  ")

        configs = {
            "일반":        ("맑은 고딕", 11, "#ffffff"),
            "마크다운":    ("Consolas",  11, "#1e1e2e"),
            "대본":        ("Courier New", 11, "#fffff8"),
            "뮤지컬 가사": ("맑은 고딕", 12, "#fffdf0"),
        }
        family, size, bg = configs.get(mode, ("맑은 고딕", 11, "#ffffff"))
        self.text_edit.setFont(QFont(family, size))
        self.text_edit.setStyleSheet(f"""
            QTextEdit {{
                background-color: {bg};
                color: #1a1a1a;
                border: none;
                font-size: {size}pt;
            }}
        """)

    # ── 파일 입출력 ───────────────────────────────────────

    def new_doc(self):
        if self.text_edit.document().isModified():
            r = QMessageBox.question(
                self, "새 문서",
                "저장하지 않은 내용이 있습니다. 계속할까요?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            if r == QMessageBox.StandardButton.No:
                return
        self.text_edit.clear()
        self._current_path = ""

    def open_doc(self):
        path, _ = QFileDialog.getOpenFileName(
            self, "파일 열기", "",
            "텍스트/마크다운 (*.txt *.md);;"
            "Word 파일 (*.docx);;"
            "모든 파일 (*)"
        )
        if path:
            self.load_file(path)

    def load_file(self, path: str):
        """사이드바 더블클릭 또는 open_doc 에서 호출"""
        ext = os.path.splitext(path)[1].lower()
        try:
            if ext in (".txt", ".md"):
                with open(path, "r", encoding="utf-8") as f:
                    self.text_edit.setPlainText(f.read())
                if ext == ".md":
                    self.mode_box.setCurrentText("마크다운")
            elif ext == ".docx":
                try:
                    from docx import Document
                    doc = Document(path)
                    text = "\n".join(p.text for p in doc.paragraphs)
                    self.text_edit.setPlainText(text)
                except ImportError:
                    QMessageBox.warning(self, "오류",
                        "python-docx 패키지가 필요합니다.\n"
                        "터미널에서: pip install python-docx")
                    return
            else:
                with open(path, "r", encoding="utf-8", errors="replace") as f:
                    self.text_edit.setPlainText(f.read())

            self._current_path = path
            self.text_edit.document().setModified(False)

        except Exception as e:
            QMessageBox.warning(self, "파일 열기 오류", str(e))

    def save_doc(self):
        if not self._current_path:
            path, _ = QFileDialog.getSaveFileName(
                self, "저장", "",
                "텍스트 파일 (*.txt);;마크다운 (*.md)"
            )
            if not path:
                return
            self._current_path = path

        try:
            with open(self._current_path, "w", encoding="utf-8") as f:
                f.write(self.text_edit.toPlainText())
            self.text_edit.document().setModified(False)
            self.saved.emit()
        except Exception as e:
            QMessageBox.warning(self, "저장 오류", str(e))

    def export_docx(self):
        path, _ = QFileDialog.getSaveFileName(
            self, "Word로 내보내기", "", "Word 파일 (*.docx)"
        )
        if not path:
            return
        try:
            from docx import Document
            from docx.shared import Pt, Cm
            doc = Document()
            section = doc.sections[0]
            section.top_margin = section.bottom_margin = Cm(2.5)
            section.left_margin = section.right_margin = Cm(3.0)
            for line in self.text_edit.toPlainText().split("\n"):
                p = doc.add_paragraph(line)
                p.style.font.size = Pt(11)
            doc.save(path)
            QMessageBox.information(self, "완료", f"저장되었습니다:\n{path}")
        except ImportError:
            QMessageBox.warning(self, "오류",
                "python-docx 패키지가 필요합니다.\n"
                "터미널에서: pip install python-docx")

    def export_pdf(self):
        path, _ = QFileDialog.getSaveFileName(
            self, "PDF로 내보내기", "", "PDF 파일 (*.pdf)"
        )
        if not path:
            return
        try:
            from weasyprint import HTML, CSS
            css = CSS(string="""
                @page { size: A4; margin: 2.5cm 3cm; }
                body {
                    font-family: 'Malgun Gothic', sans-serif;
                    font-size: 11pt;
                    line-height: 1.8;
                    color: #1a1a1a;
                }
            """)
            HTML(string=self.text_edit.toHtml()).write_pdf(path, stylesheets=[css])
            QMessageBox.information(self, "완료", f"저장되었습니다:\n{path}")
        except ImportError:
            QMessageBox.warning(self, "오류",
                "weasyprint 패키지가 필요합니다.\n"
                "터미널에서: pip install weasyprint")
